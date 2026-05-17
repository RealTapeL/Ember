import re
import tempfile
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from app.models import Tutorial, Step


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_markdown_content(doc, text: str):
    """Parse simple markdown and add to docx document."""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    lines = text.split('\n')
    in_code = False
    code_lines = []

    for line in lines:
        stripped = line.strip()

        # Code block start/end
        if stripped.startswith('```'):
            if in_code:
                # End code block
                if code_lines:
                    # Add code as table with gray background
                    table = doc.add_table(rows=1, cols=1)
                    table.style = 'Table Grid'
                    cell = table.cell(0, 0)
                    # Remove borders
                    cell._tc.get_or_add_tcPr().append(parse_xml(
                        r'<w:tcBorders {}><w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                        r'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                        r'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                        r'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/></w:tcBorders>'.format(nsdecls('w'))
                    ))
                    # Gray background
                    cell._tc.get_or_add_tcPr().append(parse_xml(
                        r'<w:shd {} w:fill="F5F5F5"/>'.format(nsdecls('w'))
                    ))
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.space_before = Pt(0)
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Courier New'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
                    run.font.size = Pt(9)
                    code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        # Heading
        if stripped.startswith('##'):
            level = len(stripped) - len(stripped.lstrip('#'))
            heading_text = stripped.lstrip('#').strip()
            doc.add_heading(heading_text, level=min(level, 3))
            continue

        # Bold inline
        if stripped.startswith('**') and stripped.endswith('**') and stripped.count('**') == 2:
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip('*'))
            run.bold = True
            run.font.size = Pt(11)
            continue

        # Bullet list
        if stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(stripped[2:], style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
            continue

        # Numbered list
        match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if match:
            p = doc.add_paragraph(match.group(2), style='List Number')
            p.paragraph_format.left_indent = Inches(0.25)
            continue

        # Empty line
        if not stripped:
            continue

        # Normal paragraph with inline formatting
        p = doc.add_paragraph()
        # Parse inline bold and italic
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', stripped)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part.strip('*'))
                run.bold = True
                run.font.size = Pt(11)
            elif part.startswith('*') and part.endswith('*'):
                run = p.add_run(part.strip('*'))
                run.italic = True
                run.font.size = Pt(11)
            else:
                run = p.add_run(part)
                run.font.size = Pt(11)


async def export_tutorial_to_docx(db: AsyncSession, tutorial_id: int) -> str:
    """Export tutorial to a Word document and return file path."""
    # Get tutorial and steps
    result = await db.execute(select(Tutorial).where(Tutorial.id == tutorial_id))
    tutorial = result.scalar_one_or_none()
    if not tutorial:
        raise ValueError("Tutorial not found")

    steps_result = await db.execute(
        select(Step).where(Step.tutorial_id == tutorial_id).order_by(Step.order)
    )
    steps = steps_result.scalars().all()

    # Create document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Title
    title = doc.add_heading(tutorial.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x2B, 0x5E)
        run.font.size = Pt(22)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Subtitle
    if tutorial.description:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(tutorial.description)
        run.italic = True
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.font.size = Pt(11)

    # Separator
    doc.add_paragraph()

    # Meta info
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('—— 矽励科技 · 知识库 自动生成 ——')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()
    doc.add_paragraph('_' * 50)
    doc.add_paragraph()

    # Steps
    for step in steps:
        # Step title
        heading = doc.add_heading(f"步骤 {step.order}：{step.title}", level=1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x2B, 0x5E)
            run.font.size = Pt(16)

        # Step description
        if step.description:
            p = doc.add_paragraph()
            run = p.add_run(step.description)
            run.italic = True
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            run.font.size = Pt(10)

        doc.add_paragraph()

        # Step content (markdown)
        if step.content:
            add_markdown_content(doc, step.content)

        # Separator between steps
        doc.add_paragraph()
        doc.add_paragraph('─' * 40)
        doc.add_paragraph()

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('本教程由 AI 基于原文档自动生成，仅供参考学习。')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Save to temp file
    fd, path = tempfile.mkstemp(suffix='.docx')
    os.close(fd)
    doc.save(path)
    return path
